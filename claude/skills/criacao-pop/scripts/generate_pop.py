#!/usr/bin/env python3
"""
Gerador de POP (Procedimento Operacional Padrão) — formato Soufit/Soulve.

Uso:
    python generate_pop.py --data pop_data.json --output POP-XXX-NNN.docx

O arquivo JSON deve seguir o schema documentado em references/pop_schema.json.

Saída: arquivo .docx no formato padronizado (tabela 3 colunas com seções mescladas).
"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm


# ============================================================
# CONFIGURAÇÕES DE FORMATAÇÃO
# ============================================================

FONT_NAME = "Calibri"
FONT_SIZE_NORMAL = 10
FONT_SIZE_HEADER = 11
HEADER_FILL = "D9D9D9"   # cinza claro nos cabeçalhos de seção
TOP_FILL = "BFBFBF"      # cinza médio no cabeçalho de identificação


# ============================================================
# HELPERS DE FORMATAÇÃO
# ============================================================

def set_cell_shading(cell, fill_hex: str):
    """Aplica cor de fundo a uma célula."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    """Garante bordas pretas finas em todos os lados da célula."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def style_cell(cell, text: str, *, bold=False, fill=None, size=FONT_SIZE_NORMAL,
               align=WD_ALIGN_PARAGRAPH.LEFT):
    """Aplica formatação padrão a uma célula."""
    cell.text = ""  # limpa default
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    if fill:
        set_cell_shading(cell, fill)
    set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_section_header(table, label: str):
    """Adiciona uma linha de cabeçalho de seção mesclando as 3 colunas."""
    row = table.add_row()
    cells = row.cells
    merged = cells[0].merge(cells[1]).merge(cells[2])
    style_cell(merged, label, bold=True, fill=HEADER_FILL, size=FONT_SIZE_HEADER,
               align=WD_ALIGN_PARAGRAPH.LEFT)


def add_merged_content_row(table, text: str):
    """Adiciona uma linha de conteúdo com as 3 colunas mescladas."""
    row = table.add_row()
    cells = row.cells
    merged = cells[0].merge(cells[1]).merge(cells[2])
    style_cell(merged, text)


def add_three_col_row(table, c1: str, c2: str, c3: str):
    """Adiciona uma linha com 3 colunas separadas."""
    row = table.add_row()
    cells = row.cells
    style_cell(cells[0], c1)
    style_cell(cells[1], c2)
    style_cell(cells[2], c3)


# ============================================================
# CONSTRUÇÃO DO DOCUMENTO
# ============================================================

def build_pop(data: dict, output_path: Path) -> None:
    """Constrói o .docx do POP a partir do dicionário de dados."""
    validate_data(data)

    doc = Document()

    # Margens A4 padrão
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Estilo padrão do documento
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE_NORMAL)

    # Título
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("Procedimento operacional padrão")
    title_run.bold = True
    title_run.font.name = FONT_NAME
    title_run.font.size = Pt(14)

    # Tabela principal — 3 colunas
    table = doc.add_table(rows=0, cols=3)
    table.autofit = False

    # ---------- Cabeçalho de identificação (3 linhas) ----------
    add_three_col_row(
        table,
        data["empresa"],
        "Procedimento operacional padrão",
        f"Padrão no {data['codigo_pop']}",
    )
    # Aplica negrito + fill no cabeçalho de identificação (linha 1)
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
        set_cell_shading(cell, TOP_FILL)

    add_three_col_row(
        table,
        f"Area: {data['area']}",
        f"Estabelecido em: {data['estabelecido_em']}",
        f"Revisado em: {data.get('revisado_em', '')}",
    )
    add_three_col_row(
        table,
        f"Nome da tarefa: {data['nome_tarefa']}",
        f"Responsavel: {data['responsavel']}",
        f"No da revisão: {data['revisao']}",
    )

    # ---------- Objetivo ----------
    add_section_header(table, "Objetivo")
    add_merged_content_row(table, format_numbered_list([data["objetivo"]]))

    # ---------- Material necessário ----------
    add_section_header(table, "Material necessário")
    for item in data["material_necessario"]:
        add_three_col_row(
            table,
            item["item"],
            str(item["quantidade"]),
            item["observacao"],
        )

    # ---------- Passos críticos ----------
    add_section_header(table, "Passos críticos")
    add_merged_content_row(table, format_numbered_list(data["passos_criticos"]))

    # ---------- Manuseio do material ----------
    add_section_header(table, "Manuseio do material")
    add_merged_content_row(table, format_numbered_list(data["manuseio_material"]))

    # ---------- Resultados esperados ----------
    add_section_header(table, "Resultados esperados")
    add_merged_content_row(table, format_numbered_list(data["resultados_esperados"]))

    # ---------- Ações corretivas ----------
    add_section_header(table, "Ações corretivas")
    add_merged_content_row(table, format_numbered_list(data["acoes_corretivas"]))

    # ---------- Aprovação ----------
    add_section_header(table, "Aprovação")
    aprovacao_text = (
        f"____________________ Executor {data.get('area_curta', data['area'])}     "
        f"____________________ Supervisor     "
        f"____________________ Chefia"
    )
    add_merged_content_row(table, aprovacao_text)

    # ---------- Fontes (fora da tabela) ----------
    doc.add_paragraph()  # espaço
    fontes_para = doc.add_paragraph()
    fontes_run = fontes_para.add_run(f"Fontes: {data['fontes']}")
    fontes_run.font.name = FONT_NAME
    fontes_run.font.size = Pt(FONT_SIZE_NORMAL)
    fontes_run.italic = True

    # Salvar
    doc.save(str(output_path))


def format_numbered_list(items: list[str]) -> str:
    """Formata uma lista de itens como '01 - texto 02 - texto ...'."""
    formatted = []
    for i, text in enumerate(items, start=1):
        num = f"{i:02d}"
        text = text.strip()
        # Garante que termina com pontuação
        if text and text[-1] not in ".!?":
            text += "."
        formatted.append(f"{num} - {text}")
    return " ".join(formatted)


# ============================================================
# VALIDAÇÃO
# ============================================================

REQUIRED_FIELDS = [
    "empresa", "codigo_pop", "area", "estabelecido_em", "nome_tarefa",
    "responsavel", "revisao", "objetivo", "material_necessario",
    "passos_criticos", "manuseio_material", "resultados_esperados",
    "acoes_corretivas", "fontes",
]


def validate_data(data: dict) -> None:
    """Valida os campos obrigatórios e regras de qualidade Falconi."""
    errors = []

    # Campos obrigatórios
    for field in REQUIRED_FIELDS:
        if field not in data:
            errors.append(f"Campo obrigatório ausente: '{field}'")

    if errors:
        raise ValueError("\n".join(errors))

    # Código do POP no formato POP-AREA-NNN
    if not data["codigo_pop"].startswith("POP-"):
        errors.append(f"codigo_pop deve começar com 'POP-' (recebido: {data['codigo_pop']})")

    # Material: pelo menos 3 itens
    if len(data["material_necessario"]) < 3:
        errors.append(
            f"material_necessario deve ter ao menos 3 itens "
            f"(recebido: {len(data['material_necessario'])})"
        )

    # Cada item de material precisa ter item/quantidade/observacao
    for i, mat in enumerate(data["material_necessario"]):
        for k in ("item", "quantidade", "observacao"):
            if k not in mat:
                errors.append(f"material_necessario[{i}] sem campo '{k}'")

    # Passos críticos: entre 10 e 30 (faixa um pouco mais flexível que o ideal 15-25)
    n_passos = len(data["passos_criticos"])
    if n_passos < 10:
        errors.append(
            f"passos_criticos tem só {n_passos} itens — POP provavelmente está raso "
            f"(mínimo recomendado: 15)"
        )
    if n_passos > 30:
        errors.append(
            f"passos_criticos tem {n_passos} itens — processo precisa ser quebrado "
            f"em mais de um POP (máximo: 30)"
        )

    # Manuseio: 4 a 8 itens
    if not (4 <= len(data["manuseio_material"]) <= 8):
        errors.append(
            f"manuseio_material deve ter entre 4 e 8 itens "
            f"(recebido: {len(data['manuseio_material'])})"
        )

    # Resultados: 5 a 10 itens
    if not (5 <= len(data["resultados_esperados"]) <= 10):
        errors.append(
            f"resultados_esperados deve ter entre 5 e 10 itens "
            f"(recebido: {len(data['resultados_esperados'])})"
        )

    # Ações corretivas: 5 a 10 itens, todos começando com "Se "
    if not (5 <= len(data["acoes_corretivas"]) <= 10):
        errors.append(
            f"acoes_corretivas deve ter entre 5 e 10 itens "
            f"(recebido: {len(data['acoes_corretivas'])})"
        )
    for i, acao in enumerate(data["acoes_corretivas"]):
        if not acao.strip().lower().startswith("se "):
            errors.append(
                f"acoes_corretivas[{i}] deve começar com 'Se ' "
                f"(recebido: '{acao[:40]}...')"
            )

    if errors:
        raise ValueError(
            "Falha na validação Falconi:\n" + "\n".join(f"  - {e}" for e in errors)
        )


# ============================================================
# CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="Gerador de POP padronizado Soufit/Soulve")
    parser.add_argument("--data", required=True, help="Caminho do arquivo JSON com dados do POP")
    parser.add_argument("--output", required=True, help="Caminho do .docx de saída")
    args = parser.parse_args()

    data_path = Path(args.data)
    output_path = Path(args.output)

    if not data_path.exists():
        print(f"❌ Arquivo de dados não encontrado: {data_path}", file=sys.stderr)
        sys.exit(1)

    with data_path.open("r", encoding="utf-8") as f:
        data = json.load(f)

    try:
        build_pop(data, output_path)
    except ValueError as e:
        print(f"❌ Validação falhou:\n{e}", file=sys.stderr)
        sys.exit(2)

    print(f"✅ POP gerado: {output_path}")
    print(f"   Código: {data['codigo_pop']}")
    print(f"   Passos críticos: {len(data['passos_criticos'])}")
    print(f"   Ações corretivas: {len(data['acoes_corretivas'])}")


if __name__ == "__main__":
    main()
