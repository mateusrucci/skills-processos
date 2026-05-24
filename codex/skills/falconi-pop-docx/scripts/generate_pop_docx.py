#!/usr/bin/env python3
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Cm, Pt
except ImportError as exc:
    raise SystemExit(
        "Missing dependency: python-docx. Install with `python3 -m pip install python-docx`."
    ) from exc


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text or ""))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_full_width_heading(table, text):
    row = table.add_row()
    row.cells[0].merge(row.cells[-1])
    set_cell_text(row.cells[0], text, bold=True)
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_numbered_lines(table, title, values):
    add_full_width_heading(table, title)
    row = table.add_row()
    row.cells[0].merge(row.cells[-1])
    lines = []
    for index, value in enumerate(values or [], start=1):
        lines.append(f"{index:02d} - {value}")
    set_cell_text(row.cells[0], "\n".join(lines) if lines else "A confirmar.")


def build_doc(data):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    title = doc.add_paragraph("Procedimento operacional padrao")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(12)

    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    row = table.add_row()
    set_cell_text(row.cells[0], data.get("empresa", "Empresa"), bold=True)
    row.cells[0].merge(row.cells[1])
    set_cell_text(row.cells[2], "Procedimento operacional padrao", bold=True)
    set_cell_text(row.cells[3], f"Padrao no {data.get('codigo', 'POP-000')}", bold=True)

    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    set_cell_text(row.cells[0], f"Area: {data.get('area', 'A confirmar')}")
    set_cell_text(row.cells[2], f"Estabelecido em: {data.get('estabelecido_em', 'A confirmar')}")
    set_cell_text(row.cells[3], f"Revisado em: {data.get('revisado_em', '')}")

    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    set_cell_text(row.cells[0], f"Nome da tarefa: {data.get('nome_tarefa', 'A confirmar')}")
    set_cell_text(row.cells[2], f"Responsavel: {data.get('responsavel', 'A confirmar')}")
    set_cell_text(row.cells[3], f"No da revisao: {data.get('numero_revisao', 'primeira')}")

    if data.get("objetivo"):
        add_numbered_lines(table, "Objetivo", [data["objetivo"]])

    add_full_width_heading(table, "Material necessario")
    materials = data.get("materiais") or []
    if materials:
        for material in materials:
            row = table.add_row()
            set_cell_text(row.cells[0], material.get("item", "A confirmar"))
            set_cell_text(row.cells[1], material.get("quantidade", ""))
            row.cells[2].merge(row.cells[3])
            set_cell_text(row.cells[2], material.get("observacao", ""))
    else:
        row = table.add_row()
        row.cells[0].merge(row.cells[-1])
        set_cell_text(row.cells[0], "A confirmar.")

    add_numbered_lines(table, "Passos criticos", data.get("passos_criticos"))
    add_numbered_lines(table, "Manuseio do material", data.get("manuseio_material"))
    add_numbered_lines(table, "Resultados esperados", data.get("resultados_esperados"))
    add_numbered_lines(table, "Acoes corretivas", data.get("acoes_corretivas"))

    add_full_width_heading(table, "Aprovacao")
    row = table.add_row()
    approvers = data.get("aprovadores") or ["Executor", "Supervisor", "Chefia"]
    approval_text = "     ".join([f"____________________\n{role}" for role in approvers])
    row.cells[0].merge(row.cells[-1])
    set_cell_text(row.cells[0], approval_text)
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if data.get("fontes"):
        doc.add_paragraph()
        source = doc.add_paragraph("Fontes: " + "; ".join(data["fontes"]))
        source.runs[0].font.size = Pt(8)

    return doc


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: generate_pop_docx.py input.json output.docx")

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    data = json.loads(input_path.read_text(encoding="utf-8"))
    doc = build_doc(data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Created {output_path}")


if __name__ == "__main__":
    main()
